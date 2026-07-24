import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generer_rapport_word(contexte, resultats_par_etat, texte_interpretation=None, figure_combinee_bytes=None):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    title = doc.add_heading("RAPPORT DE PREVISION FINANCIAL ANALYTICS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(f"Bank of Africa — Direction de la Trésorerie & des Moyens de Paiement\nMoyen : {contexte['moyen']} | Variable : {contexte['variable']}")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading("1. Périmètre & Contexte de l'Étude", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Établissement Bancaire : ").bold = True
    p.add_run(f"{contexte['banque']}\n")
    p.add_run(f"Moyen de Paiement : ").bold = True
    p.add_run(f"{contexte['moyen']}\n")
    p.add_run(f"Indicateur Modélisé : ").bold = True
    p.add_run(f"{contexte['variable']}\n")
    p.add_run(f"Horizon Temporel : ").bold = True
    p.add_run(f"{contexte['horizon_annees']} an(s)\n\n")
    
    p.add_run("Cadre Méthodologique : ").bold = True
    p.add_run("L'analyse combine l'historique observable avec un prolongement calibré intégrant la tendance séculaire, la saisonnalité sectorielle et le calendrier lunaire (Ramadan/Aïd). La sélection des modèles s'effectue par validation croisée sur le sMAPE out-of-sample.")

    for etat, res in resultats_par_etat.items():
        nom_etat = "Flux Émis" if etat == "Emis" else "Flux Reçus"
        doc.add_heading(f"2. Modélisation Économétrique — {nom_etat}", level=1)
        if res.get("erreur"):
            doc.add_paragraph(f"Notification : {res['erreur']}")
            continue
        
        p_res = doc.add_paragraph()
        p_res.add_run(f"Spécification SARIMA retenue : ").bold = True
        p_res.add_run(f"SARIMA{res['meilleur']['ordre']}\n")
        p_res.add_run(f"Akaike Info Criterion (AIC) : ").bold = True
        p_res.add_run(f"{res['meilleur']['AIC']}\n")
        p_res.add_run(f"Précision Test sMAPE : ").bold = True
        p_res.add_run(f"{res['meilleur']['sMAPE_%']}%\n")
        
        if "figure_bytes" in res:
            doc.add_picture(io.BytesIO(res["figure_bytes"]), width=Inches(6.2))

    if len(resultats_par_etat) == 2 and all(not r.get("erreur") for r in resultats_par_etat.values()):
        doc.add_heading("3. Analyse Comparative & Position Nette", level=1)
        emis, recus = resultats_par_etat.get("Emis", {}), resultats_par_etat.get("Recus", {})
        
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Métriques Économétriques", "Flux Émis", "Flux Reçus"
        table.rows[1].cells[0].text = "Modèle Optima SARIMA"
        table.rows[1].cells[1].text = emis.get("meilleur", {}).get("ordre", "-")
        table.rows[1].cells[2].text = recus.get("meilleur", {}).get("ordre", "-")
        table.rows[2].cells[0].text = "Erreur sMAPE (Test %)"
        table.rows[2].cells[1].text = str(emis.get("meilleur", {}).get("sMAPE_%", "-"))
        table.rows[2].cells[2].text = str(recus.get("meilleur", {}).get("sMAPE_%", "-"))

        if figure_combinee_bytes:
            doc.add_paragraph()
            doc.add_picture(io.BytesIO(figure_combinee_bytes), width=Inches(6.2))

        if texte_interpretation:
            doc.add_heading("Synthèse Analytique & Recommandations", level=2)
            texte_clean = texte_interpretation.replace("#", "").replace("**", "")
            doc.add_paragraph(texte_clean)

    doc.add_heading("4. Clarifications Méthodologiques & Limites", level=1)
    doc.add_paragraph(
        "1. L'extension synthétique constitue un outil de simulation robuste calibré sur les lois statistiques historiques.\n"
        "2. Les coefficients d'impact liés au calendrier lunaire (Ramadan/Aïd) sont réestimés continuellement selon l'historique effectif.\n"
        "3. Les prévisions doivent faire l'objet de révisions périodiques à chaque clôture mensuelle."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer